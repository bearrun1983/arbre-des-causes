# fichier: arbre_des_causes_app.py
import os
import json
from io import BytesIO
from collections import defaultdict, deque

import streamlit as st
import graphviz
from docx import Document  # python-docx

# =============== CONSTANTES ===============
CATEGORIES = {
    "ORGANISATIONNELLE": {"color": "#5B9BD5", "desc": "Bleu soutenu"},
    "HUMAINE": {"color": "#ED7D31", "desc": "Orange soutenu"},
    "TECHNIQUE": {"color": "#A6A6A6", "desc": "Gris soutenu"},
}
RANKDIR = "RL"  # racine à droite -> causes à gauche
ARROW_MODE = "PARENT_TO_CHILD"  # flèches Parent -> Enfant

# =============== ETATS INITIAUX ===============
if "page" not in st.session_state:
    st.session_state.page = "Accueil"

# Arbre des causes
if "nodes" not in st.session_state:
    st.session_state.nodes = {"root": {"label": "Racine", "category": None}}
if "edges" not in st.session_state:
    st.session_state.edges = []
if "root_label" not in st.session_state:
    st.session_state.root_label = "Racine"

# 5 Pourquoi
if "why" not in st.session_state:
    st.session_state.why = []
if "why_problem" not in st.session_state:
    st.session_state.why_problem = ""

# Assistant IA
if "ai_doc_text" not in st.session_state:
    st.session_state.ai_doc_text = ""
if "ai_struct" not in st.session_state:
    st.session_state.ai_struct = None  # dictionnaire structuré (questions/faits/causes)
if "ai_questions_flat" not in st.session_state:
    st.session_state.ai_questions_flat = []  # liste plate (pour export)

# =============== HELPERS GLOBAUX ===============
def get_parent(node_id: str):
    for src, tgt in st.session_state.edges:
        if tgt == node_id:
            return src
    return None

def build_children_map(edges):
    children = defaultdict(list)
    for src, tgt in edges:
        children[src].append(tgt)
    return children

def is_descendant(root_id: str, query_id: str) -> bool:
    """Retourne True si query_id est dans le sous-arbre de root_id (Parent->Enfant)."""
    children = build_children_map(st.session_state.edges)
    q = deque([root_id])
    while q:
        n = q.popleft()
        if n == query_id:
            return True
        q.extend(children.get(n, []))
    return False

def export_arbre_docx(title, nodes, edges) -> BytesIO:
    doc = Document()
    doc.add_heading(title or "Arbre des causes", 0)

    doc.add_heading("Catégories", level=1)
    for cat, info in CATEGORIES.items():
        doc.add_paragraph(f"- {cat} : {info['desc']}")

    doc.add_heading("Nœuds", level=1)
    for nid, data in nodes.items():
        label = data.get("label", "")
        cat = data.get("category") or "Non défini"
        doc.add_paragraph(f"- {label} ({cat})")

    doc.add_heading("Liens Parent → Enfant", level=1)
    for src, tgt in edges:
        doc.add_paragraph(f"{nodes[src]['label']} → {nodes[tgt]['label']}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_why_docx(problem, answers) -> BytesIO:
    doc = Document()
    doc.add_heading("Analyse 5 Pourquoi", 0)
    if problem:
        doc.add_paragraph(f"Problème observé : {problem}")
    for i, ans in enumerate(answers, 1):
        if ans.strip():
            doc.add_paragraph(f"{i}. Pourquoi ? — {ans.strip()}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_questions_docx(problem, ai_struct) -> BytesIO:
    """Export des questions IA (par thèmes) + faits + pistes causes."""
    doc = Document()
    doc.add_heading("Questions d’enquête (Assistant IA)", 0)
    if problem:
        doc.add_paragraph(f"Problème observé : {problem}")

    if ai_struct and "questions" in ai_struct:
        for theme, items in ai_struct["questions"].items():
            doc.add_heading(f"Thème : {theme}", level=1)
            for it in items:
                q = it.get("question", "")
                why = it.get("rationale", "")
                proof = it.get("evidence", "")
                p = doc.add_paragraph(style=None)
                p.add_run("• ").bold = True
                p.add_run(q)
                if why:
                    doc.add_paragraph(f"  - Pourquoi : {why}")
                if proof:
                    doc.add_paragraph(f"  - Preuves attendues : {proof}")

    if ai_struct and ai_struct.get("facts"):
        doc.add_heading("Faits identifiés", level=1)
        for f in ai_struct["facts"]:
            doc.add_paragraph(f"• {f}")

    if ai_struct and ai_struct.get("candidate_causes"):
        doc.add_heading("Pistes de causes (à investiguer)", level=1)
        for c in ai_struct["candidate_causes"]:
            lbl = c.get("label", "")
            cat = c.get("category", "")
            just = c.get("justification", "")
            doc.add_paragraph(f"• {lbl} [{cat}]")
            if just:
                doc.add_paragraph(f"  - Justification : {just}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# -------- Extraction texte .docx (paragraphes + tableaux) --------
def extract_docx_text(file_bytes: BytesIO) -> str:
    doc = Document(file_bytes)
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # tableaux -> cellules concaténées par tabulation, lignes par saut de ligne
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(c for c in cells if c))
    return "\n".join(parts).strip()

# -------- Utilitaires JSON --------
def try_extract_json(text: str):
    """Extrait le premier bloc JSON plausible dans un texte."""
    if not text:
        return None
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        snippet = text[start : end + 1]
        try:
            return json.loads(snippet)
        except Exception:
            return None
    return None

def normalize_category(s: str):
    if not s:
        return None
    s = s.strip().lower()
    if "huma" in s:
        return "HUMAINE"
    if "orga" in s:
        return "ORGANISATIONNELLE"
    if "tech" in s:
        return "TECHNIQUE"
    return None

# -------- IA: OpenAI (JSON strict) ou heuristique locale --------
def ai_analyze_text(text: str):
    """
    Retourne un dict structuré :
    {
      "summary": "...",
      "questions": {
         "Chronologie": [{"question": "...", "rationale":"...", "evidence":"..."}],
         "Organisation": [...],
         "Humain": [...],
         "Technique": [...],
         "Environnement": [...],
         "Barrières/Contrôles": [...]
      },
      "facts": ["...", "..."],
      "candidate_causes": [{"label":"...", "category":"...", "justification":"..."}]
    }
    """
    api_key = os.environ.get("OPENAI_API_KEY")
    if api_key:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)

            # Prompt orienté HSE, sortie JSON strict
            system = (
                "Tu es un expert HSE qui prépare une interview d'enquête après accident. "
                "Tu DOIS répondre UNIQUEMENT en JSON valide selon le schéma donné."
            )
            user = f"""
Analyse le recueil d'effets ci-dessous (accident du travail) et PRODUIS des QUESTIONS d'enquête
vraiment utiles (ouvertes, précises, orientées action), avec une brève justification et des preuves attendues.
Propose aussi des FAITS objectifs (formulations brèves, neutres) et des PISTES DE CAUSES à investiguer
avec une catégorie proposée (HUMAINE / ORGANISATIONNELLE / TECHNIQUE) et une justification.

Recueil d'effets:
\"\"\"{text}\"\"\"

Réponds en JSON strict suivant CE SCHÉMA EXACT (clés et structure) :
{{
  "summary": "2-4 lignes résumant l'événement et les zones d'incertitude",
  "questions": {{
    "Chronologie": [{{"question":"...", "rationale":"...", "evidence":"..."}}],
    "Organisation": [{{"question":"...", "rationale":"...", "evidence":"..."}}],
    "Humain": [{{"question":"...", "rationale":"...", "evidence":"..."}}],
    "Technique": [{{"question":"...", "rationale":"...", "evidence":"..."}}],
    "Environnement": [{{"question":"...", "rationale":"...", "evidence":"..."}}],
    "Barrières/Contrôles": [{{"question":"...", "rationale":"...", "evidence":"..."}}]
  }},
  "facts": ["...", "..."],
  "candidate_causes": [{{"label":"...", "category":"HUMAINE|ORGANISATIONNELLE|TECHNIQUE", "justification":"..."}}]
}}
- Les QUESTIONS doivent aider à faire émerger des informations manquantes, contradictions, décisions, signaux faibles.
- "evidence" = documents/observations/preuves pratiques à collecter (FDS, permis, consignation, radios, planning, météo, photos, entretiens...).
- N'invente pas de faits ; si une info manque, transforme-la en question.
"""

            # Utilisation chat.completions (compat simple) ; on impose JSON-only via prompt
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "system", "content": system},
                          {"role": "user", "content": user}],
                temperature=0.2,
            )
            content = resp.choices[0].message.content
            data = try_extract_json(content)
            if not data:
                # fallback minimal si parsing échoue
                data = heuristic_struct(text)
            # Normaliser catégories
            for c in data.get("candidate_causes", []):
                c["category"] = normalize_category(c.get("category")) or c.get("category")
            return data
        except Exception as e:
            st.warning(f"IA OpenAI indisponible ({e}). Passage en mode local.")
            return heuristic_struct(text)
    else:
        return heuristic_struct(text)

def heuristic_struct(text: str):
    """Heuristique locale sérieuse : questions par thèmes + faits + pistes de causes."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    # détecter quelques mots-clés
    lower = ("\n".join(lines)).lower()

    # Faits: phrases courtes/punaisées
    facts = []
    for l in lines:
        if (l.startswith(("-", "*", "•")) or len(l.split()) <= 12) and ("?" not in l):
            facts.append(l.lstrip("-*• ").strip(". "))

    if not facts:
        facts = ["Faits à confirmer : lieu/heure exacte, tâche en cours, équipements impliqués."]

    def Q(q, r, e):  # mini helper
        return {"question": q, "rationale": r, "evidence": e}

    questions = {
        "Chronologie": [
            Q("Déroulez minute par minute les 60 minutes précédant l’événement.",
              "Identifier enchaînement réel vs prévu, écarts et déclencheurs.",
              "Feuilles de route, main courante, enregistrements radio/téléphone, badges."),
            Q("Quelles décisions ou changements de plan ont eu lieu le jour J ? Par qui, pourquoi ?",
              "Repérer les arbitrages et contraintes réelles.",
              "Briefing d'équipe, mails, ordres de travail, main courante."),
        ],
        "Organisation": [
            Q("Quelles procédures/consignations/permits étaient applicables ? Ont-elles été comprises et suivies ?",
              "Mettre en évidence l’adéquation procédure–terrain.",
              "Procédures, permis, signatures, formation, entretiens."),
            Q("La charge de travail, l'effectif et la supervision étaient-ils adaptés ?",
              "Détecter sous-staffing, multitâche, pression délai.",
              "Planning, affectations, main courante, entretiens superviseur."),
        ],
        "Humain": [
            Q("Quelles compétences/expériences spécifiques avaient les intervenants pour cette tâche ?",
              "Vérifier adéquation compétence–risque.",
              "Registres de formation, habilitations, entretiens."),
            Q("Fatigue, stress, distraction : des signaux étaient-ils présents ?",
              "Explorer facteurs humains non documentés.",
              "Entretiens, horaires, pauses, charge mentale rapportée."),
        ],
        "Technique": [
            Q("Quel était l’état réel des équipements (défauts connus, maintenance en cours, interlocks actifs) ?",
              "Comprendre la défaillance matérielle potentielle.",
              "Historique GMAO, fiches maintenance, rapports d’essai, photos."),
            Q("Quels EPI/protections mécaniques étaient requis et effectivement portés/en place ?",
              "Tester la dernière barrière.",
              "Consignes EPI, inventaire, photos, témoignages."),
        ],
        "Environnement": [
            Q("Conditions météo/visibilité/bruit/éclairage : en quoi ont-elles influencé l’événement ?",
              "Facteurs contextuels souvent déterminants.",
              "Données météo, luxmètre, photos, mesures bruit."),
            Q("Y avait-il d'autres activités à proximité générant des interférences ?",
              "Risque d’interactions non prévues.",
              "Planning global, permis simultanés, main courante."),
        ],
        "Barrières/Contrôles": [
            Q("Quelles barrières de prévention/protection étaient prévues ? Laquelle a échoué en premier ?",
              "Situer la première défaillance barrière.",
              "Analyse bow-tie, matrices risques, procédures."),
            Q("Quels contrôles de dernière minute (point d’arrêt, LOTO, check-list) ont été réalisés ?",
              "Vérifier l’exécution des ‘dernier regard’ critiques.",
              "Check-lists signées, témoignages, enregistrements."),
        ],
    }

    # Domaines spécifiques autoroutiers si indices présents
    if any(k in lower for k in ["autorout", "circulation", "balisage", "signalisation", "trafic"]):
        questions["Environnement"].append(
            Q("Le balisage/ITPC et la gestion du trafic étaient-ils conformes (espacements, limitations, visibilité) ?",
              "Sécurité en milieu circulant = barrière majeure.",
              "Plan de balisage, photos, chrono, enregistrements trafic.")
        )
        questions["Organisation"].append(
            Q("Les communications radio avec PC trafic/ASTREINTE ont-elles couvert les phases clés ?",
              "Coordination multi-acteurs critique.",
              "Logs radio, scripts, attestations.")
        )

    candidate_causes = [
        {"label": "Procédure inadaptée au terrain réel", "category": "ORGANISATIONNELLE",
         "justification": "Écarts entre décrit et faisable, décisions ad hoc observées."},
        {"label": "Formation/habilitation insuffisante pour la tâche spécifique", "category": "HUMAINE",
         "justification": "Compétences non alignées avec le niveau de risque."},
        {"label": "Défaillance matérielle non détectée", "category": "TECHNIQUE",
         "justification": "Historique maintenance/inspection à vérifier."},
    ]

    return {
        "summary": "Synthèse à préciser à partir des éléments recueillis.",
        "questions": questions,
        "facts": facts[:30],
        "candidate_causes": candidate_causes,
    }

# =============== NAVIGATION ===============
st.sidebar.title("Navigation")
page = st.sidebar.radio(
    "Aller vers :",
    ["Accueil", "Arbre des causes", "5 Pourquoi", "Assistant IA (Recueil d’effets)"],
    index=["Accueil", "Arbre des causes", "5 Pourquoi", "Assistant IA (Recueil d’effets)"].index(st.session_state.page)
)
st.session_state.page = page

# =============== PAGES ===============
if page == "Accueil":
    st.title("Bienvenue 👋")
    st.markdown(
        """
- **Arbre des causes** : cartographier les causes multiples (catégories, édition, export Word).
- **5 Pourquoi** : remonter linéairement à la cause racine (export Word).
- **Assistant IA (Recueil d’effets)** : uploadez un **.docx** ; l’IA propose des **questions d’enquête profondes**, des **faits** et des **pistes de causes** ; vous pouvez **injecter** les faits/causes dans l’arbre.
        """
    )

# ---------- ARBRE DES CAUSES ----------
elif page == "Arbre des causes":
    st.title("Analyse par Arbre des causes")

    # Nom racine
    st.subheader("Nom de la racine")
    st.session_state.root_label = st.text_input("Nom de la racine", value=st.session_state.root_label)
    st.session_state.nodes["root"]["label"] = st.session_state.root_label

    # Ajouter un nœud
    st.subheader("Ajouter un nœud")
    new_node_label = st.text_input("Nom du nouveau nœud", key="add_label")
    parent_id = st.selectbox(
        "Sélectionner le nœud parent",
        options=list(st.session_state.nodes.keys()),
        format_func=lambda x: st.session_state.nodes[x]["label"],
        key="add_parent"
    )
    new_node_category = st.selectbox("Catégorie", options=list(CATEGORIES.keys()), index=0, key="add_cat")
    if st.button("Ajouter", key="add_btn"):
        if new_node_label.strip():
            new_node_id = f"node_{len(st.session_state.nodes)}"
            st.session_state.nodes[new_node_id] = {"label": new_node_label.strip(), "category": new_node_category}
            st.session_state.edges.append((parent_id, new_node_id))
            st.success(f"Nœud ajouté : {new_node_label}")
        else:
            st.warning("Veuillez entrer un nom de nœud valide.")

    # Modifier un nœud
    st.subheader("Modifier un nœud existant")
    node_to_edit = st.selectbox(
        "Choisir le nœud à modifier",
        options=list(st.session_state.nodes.keys()),
        format_func=lambda x: st.session_state.nodes[x]["label"],
        key="edit_select"
    )
    cur_label = st.session_state.nodes[node_to_edit]["label"]
    cur_cat = st.session_state.nodes[node_to_edit].get("category")
    cur_parent = get_parent(node_to_edit)
    edit_label = st.text_input("Nouveau libellé", value=cur_label, key="edit_label")
    edit_cat_index = list(CATEGORIES.keys()).index(cur_cat) if cur_cat in CATEGORIES else 0
    edit_cat = st.selectbox("Nouvelle catégorie", options=list(CATEGORIES.keys()), index=edit_cat_index, key="edit_cat")

    # Parent candidates sans cycles
    parents_candidates = [nid for nid in st.session_state.nodes.keys() if nid != node_to_edit]
    parents_candidates = [nid for nid in parents_candidates if not is_descendant(node_to_edit, nid)]

    if node_to_edit == "root":
        st.info("La racine ne peut pas être rattachée à un parent.")
        edit_parent = None
    else:
        default_parent_idx = parents_candidates.index(cur_parent) if cur_parent in parents_candidates else 0
        edit_parent = st.selectbox(
            "Nouveau parent",
            options=parents_candidates,
            index=default_parent_idx if parents_candidates else 0,
            format_func=lambda x: st.session_state.nodes[x]["label"],
            key="edit_parent"
        ) if parents_candidates else None

    if st.button("Mettre à jour", key="edit_btn"):
        st.session_state.nodes[node_to_edit]["label"] = edit_label.strip() or cur_label
        st.session_state.nodes[node_to_edit]["category"] = edit_cat
        if node_to_edit != "root" and edit_parent is not None and edit_parent != cur_parent:
            st.session_state.edges = [(s, t) for (s, t) in st.session_state.edges if not (s == cur_parent and t == node_to_edit)]
            st.session_state.edges.append((edit_parent, node_to_edit))
        st.success("Nœud mis à jour.")
        st.rerun()

    # Visualisation
    st.subheader("Visualisation")
    dot = graphviz.Digraph("Arbre des Causes", format="png")
    dot.attr(rankdir=RANKDIR)
    for node_id, data in st.session_state.nodes.items():
        label = data.get("label", node_id)
        cat = data.get("category")
        if cat in CATEGORIES:
            dot.node(node_id, label, style="filled", fillcolor=CATEGORIES[cat]["color"])
        else:
            dot.node(node_id, label)
    for src, tgt in st.session_state.edges:
        if ARROW_MODE == "PARENT_TO_CHILD":
            dot.edge(src, tgt)
        else:
            dot.edge(tgt, src)
    st.graphviz_chart(dot)

    # Export Word
    st.subheader("Exporter")
    if st.button("Exporter en Word (.docx)", key="export_arbre"):
        buf = export_arbre_docx(st.session_state.root_label, st.session_state.nodes, st.session_state.edges)
        st.download_button(
            "Télécharger le fichier Word",
            buf,
            file_name="arbre_des_causes.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# ---------- 5 POURQUOI ----------
elif page == "5 Pourquoi":
    st.title("Analyse par la méthode des 5 Pourquoi")
    st.markdown("**Conseils : viser 5 pourquoi.**")

    st.session_state.why_problem = st.text_area("1) Décrire le problème", value=st.session_state.why_problem)
    st.subheader("2) Poser successivement 'Pourquoi ?'")

    col_a, col_b, col_c = st.columns([1,1,1])
    with col_a:
        if st.button("Ajouter un 'Pourquoi'"):
            st.session_state.why.append("")
    with col_b:
        if st.button("Retirer le dernier"):
            if st.session_state.why:
                st.session_state.why.pop()
    with col_c:
        if st.button("Réinitialiser"):
            st.session_state.why = []
            st.rerun()

    for i in range(len(st.session_state.why)):
        st.session_state.why[i] = st.text_input(f"Réponse au Pourquoi n°{i+1}", value=st.session_state.why[i], key=f"why_{i}")

    st.subheader("Récapitulatif")
    if st.session_state.why_problem:
        st.write(f"**Problème observé :** {st.session_state.why_problem}")
    for i, ans in enumerate(st.session_state.why, 1):
        if ans.strip():
            st.write(f"{i}. Pourquoi ? — {ans.strip()}")

    st.subheader("Exporter")
    if st.button("Exporter en Word (.docx)", key="export_why"):
        buf = export_why_docx(st.session_state.why_problem, st.session_state.why)
        st.download_button(
            "Télécharger le fichier Word",
            buf,
            file_name="analyse_5_pourquoi.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# ---------- ASSISTANT IA ----------
elif page == "Assistant IA (Recueil d’effets)":
    st.title("Assistant IA (Recueil d’effets)")
    st.markdown("Uploadez votre **fichier Word (.docx)** : l’IA proposera des **questions d’enquête profondes**, des **faits** et des **pistes de causes**.")

    up = st.file_uploader("Importer un fichier Word (.docx)", type=["docx"])
    if up is not None:
        file_bytes = BytesIO(up.read())
        try:
            st.session_state.ai_doc_text = extract_docx_text(file_bytes)
            st.success("Texte extrait du .docx.")
        except Exception as e:
            st.error(f"Impossible de lire le .docx : {e}")
            st.session_state.ai_doc_text = ""

    if st.session_state.ai_doc_text:
        with st.expander("Texte extrait (aperçu)"):
            st.text_area("Texte", value=st.session_state.ai_doc_text, height=220, label_visibility="collapsed")

        if st.button("Analyser avec IA", key="ai_analyze"):
            data = ai_analyze_text(st.session_state.ai_doc_text)
            st.session_state.ai_struct = data
            # liste plate de questions (pour export/copier)
            flat = []
            for theme, items in data.get("questions", {}).items():
                for it in items:
                    flat.append(f"[{theme}] {it.get('question','')}")
            st.session_state.ai_questions_flat = flat
            st.success("Analyse IA terminée.")

    data = st.session_state.ai_struct
    if data:
        if data.get("summary"):
            st.subheader("Résumé IA")
            st.write(data["summary"])

        st.subheader("Questions par thème")
        for theme, items in data.get("questions", {}).items():
            with st.expander(f"{theme} ({len(items)})", expanded=False):
                for it in items:
                    st.write(f"**Q :** {it.get('question','')}")
                    if it.get("rationale"):
                        st.caption(f"Pourquoi : {it['rationale']}")
                    if it.get("evidence"):
                        st.caption(f"Preuves attendues : {it['evidence']}")
                    st.markdown("---")

        st.subheader("Faits proposés (sélection pour injection dans l’Arbre)")
        selected_facts = []
        for i, fact in enumerate(data.get("facts", [])):
            if st.checkbox(fact, key=f"fact_{i}"):
                selected_facts.append(fact)

        st.subheader("Pistes de causes (sélection pour injection dans l’Arbre)")
        selected_causes = []
        for i, c in enumerate(data.get("candidate_causes", [])):
            label = c.get("label", "")
            cat = normalize_category(c.get("category")) or c.get("category") or "TECHNIQUE"
            just = c.get("justification", "")
            txt = f"{label} [{cat}] — {just}"
            if st.checkbox(txt, key=f"cause_{i}"):
                selected_causes.append({"label": label, "category": cat})

        if selected_facts or selected_causes:
            st.markdown("**Paramétrage d'injection**")
            inj_parent = st.selectbox(
                "Parent",
                options=list(st.session_state.nodes.keys()),
                format_func=lambda x: st.session_state.nodes[x]["label"],
                key="inj_parent_ai"
            )
            inj_cat_default = "ORGANISATIONNELLE"
            inj_cat = st.selectbox("Catégorie par défaut (pour les faits)", options=list(CATEGORIES.keys()),
                                   index=list(CATEGORIES.keys()).index(inj_cat_default), key="inj_cat_ai")

            if st.button("Ajouter à l’Arbre", key="inj_btn_ai"):
                count = 0
                # Faits (tous utilisent la catégorie choisie)
                for fact in selected_facts:
                    new_id = f"node_{len(st.session_state.nodes)}"
                    st.session_state.nodes[new_id] = {"label": fact, "category": inj_cat}
                    st.session_state.edges.append((inj_parent, new_id))
                    count += 1
                # Causes candidates (gardent leur catégorie suggérée)
                for c in selected_causes:
                    new_id = f"node_{len(st.session_state.nodes)}"
                    st.session_state.nodes[new_id] = {"label": c["label"], "category": c["category"]}
                    st.session_state.edges.append((inj_parent, new_id))
                    count += 1
                st.success(f"{count} élément(s) ajouté(s) à l’Arbre.")
                # décocher
                for i in range(len(data.get("facts", []))):
                    if f"fact_{i}" in st.session_state:
                        st.session_state[f"fact_{i}"] = False
                for i in range(len(data.get("candidate_causes", []))):
                    if f"cause_{i}" in st.session_state:
                        st.session_state[f"cause_{i}"] = False
                st.rerun()

        st.subheader("Exporter les questions (Word)")
        if st.button("Exporter le pack de questions (.docx)", key="exp_q_docx"):
            buf = export_questions_docx(
                st.session_state.why_problem or st.session_state.root_label, data
            )
            st.download_button(
                "Télécharger les questions (Word)",
                buf,
                file_name="questions_enquete_IA.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    # Aide sur la clé
    if not os.environ.get("OPENAI_API_KEY"):
        st.info("💡 Pour activer l’IA OpenAI (meilleures questions), ajoute une clé **OPENAI_API_KEY** dans les *Secrets* Streamlit ou les variables d’environnement.")
